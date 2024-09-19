import utils
import picture

import docx
import math
import streamlit as st
import datetime
from docxtpl import DocxTemplate
from io import BytesIO
import zipfile
import copy
from lxml import etree

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.oxml import register_element_cls
from docx.enum.table import WD_TABLE_ALIGNMENT
from dataclasses import dataclass

NAME_KEY = "student_name"
CERTIFICATE_KEY = "certificate_number"
MACHINE_CATEGORY = "machine_category"

TRACTOR_PROFESSION_WORDING = "19203 «Тракторист»"

CERT_HEIGHT_INCHES = Inches(3.65)
CERT_WIDTH_INCHES = Inches(5.6)

TRACTOR_CERT_HEIGHT = Inches(5.63)
TRACTOR_CERT_WIDTH = Inches(8.04)

register_element_cls("wp:anchor", picture.CT_Anchor)


def create_confirmation_page(replacement_dict, students, picture_path):
    if not students:
        return Document()

    merged_doc = Document()
    merged_doc = utils.fit_more_rows(merged_doc)
    utils.set_default_font(merged_doc)

    merged_table = merged_doc.add_table(rows=len(students), cols=2)
    merged_tractor_table = merged_doc.add_table(rows=len(students), cols=2)

    curr_index = 0
    for student_index, student in enumerate(students):
        local_dict = replacement_dict.copy()
        local_dict[NAME_KEY] = student[NAME_KEY]
        local_dict[CERTIFICATE_KEY] = student[CERTIFICATE_KEY]
        local_dict[MACHINE_CATEGORY] = student[MACHINE_CATEGORY]

        doc = DocxTemplate("templates/milana_conf_page.docx")
        doc.render(local_dict)

        add_student_content_to_merged_table(
            merged_table, doc.tables[0], student_index, curr_index, picture_path, picture_height=Inches(5.54), picture_width=Inches(7.85)
        )
        add_student_content_to_merged_table(
            merged_tractor_table, doc.tables[1], student_index, curr_index, picture_path, picture_height=Inches(5.54), picture_width=Inches(7.85)
        ) 
        curr_index += 1

    return merged_doc

def copy_text_and_formatting(source_cell, target_cell):
    utils.copy_cell_properties(source_cell, target_cell)

    for p_i, paragraph in enumerate(source_cell.paragraphs):
        if paragraph.text.strip() == "": 
            continue
        if p_i == 0: 
            new_paragraph = target_cell.paragraphs[0]  # Use the existing empty paragraph
        else:
            new_paragraph = target_cell.add_paragraph()
        new_paragraph.paragraph_format.space_before = Pt(0)
        new_paragraph.paragraph_format.space_after = Pt(0)
        new_paragraph.alignment = paragraph.alignment
        new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent

        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            if "prof_educ_logo" in run.text:
                new_run.text = new_run.text.replace("prof_educ_logo", "") 
                new_run.add_picture('pictures/professional-education-logo.png')
                continue
            utils.preserve_formatting(new_run, run) 

def maybe_add_nested_table(cell, target_cell): 
    if len(cell.tables) > 0: 
        # for now just copy the first one
        # Set "Spacing After" for the last paragraph to 0
        last_paragraph = target_cell.paragraphs[-1]
        last_paragraph.paragraph_format.space_after = Pt(0) 
        
        nested_table = cell.tables[0] 
        new_table = target_cell.add_table(rows=len(nested_table.rows), cols=len(nested_table.columns))
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER 
        for r_i, rw in enumerate(nested_table.rows):
            for c_i, cll in enumerate(rw.cells):
                copy_text_and_formatting(cll, new_table.cell(r_i, c_i))

def add_table(merged_table, curr_row, curr_col, table): 
    for row_index, row in enumerate(table.rows):
        merged_table.rows[curr_row].height = Inches(2.76)
        for col_index, cell in enumerate(row.cells):
            target_cell = merged_table.cell(curr_row, curr_col)
            target_cell.width = Inches(3.84)
            # add tables too! 
            copy_text_and_formatting(cell, target_cell)
            maybe_add_nested_table(cell, target_cell)
            

def create_certificate_for_labour_protection(replacement_dict, students):
    if not students:
        return Document()

    merged_doc = Document()
    merged_doc = utils.fit_more_rows(merged_doc)
    utils.set_default_font(merged_doc)

    num_rows = math.ceil(len(students) / 2)

    merged_table_front = merged_doc.add_table(rows=num_rows, cols=2)
    merged_table_back = merged_doc.add_table(rows=num_rows, cols=2)


    curr_row = 0
    curr_col = 0 
    for student_index, student in enumerate(students):
        local_dict = replacement_dict.copy()
        local_dict[NAME_KEY] = student[NAME_KEY]
        local_dict[CERTIFICATE_KEY] = student[CERTIFICATE_KEY]
        local_dict[MACHINE_CATEGORY] = student[MACHINE_CATEGORY]

        doc = DocxTemplate("templates/labour_protection.docx")
        doc.render(local_dict)

        # Copy content from the template document to the target cell
        
        add_table(merged_table_front, curr_row, curr_col, doc.tables[0])
        add_table(merged_table_back, curr_row, curr_col, doc.tables[1])
           
        # Update cell indices for the next student
        curr_col += 1 
        if curr_col == 2:  
            curr_col = 0
            curr_row += 1 
    xml = etree.tostring(merged_table_front._tbl, encoding='unicode', pretty_print=True)
    utils.dump(xml, 'table')
    return merged_doc

def create_certificate(replacement_dict, students):
    if not students:
        return Document()

    all_paragraphs = []
    for student in students[1:]:
        local_dict = replacement_dict.copy()
        local_dict[NAME_KEY] = student[NAME_KEY]
        local_dict[CERTIFICATE_KEY] = student[CERTIFICATE_KEY]

        doc = DocxTemplate("templates/свидетельство.docx")
        doc.render(local_dict)

        paragraphs = doc.tables[0].cell(0, 0).paragraphs
        all_paragraphs.append(paragraphs)

    # Create final document using the first student's data as a base
    final_doc = DocxTemplate("templates/свидетельство.docx")
    local_dict = replacement_dict.copy()
    local_dict[NAME_KEY] = students[0][NAME_KEY]
    local_dict[CERTIFICATE_KEY] = students[0][CERTIFICATE_KEY]
    final_doc.render(local_dict)

    # Set default font style
    utils.set_default_font(final_doc, bold=True)

    # Get the table and add rows for each additional student
    table = final_doc.tables[0]
    for paragraphs in all_paragraphs:
        row = table.add_row()
        # this ensure that the rows are not split between pages https://github.com/python-openxml/python-docx/issues/245
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))

        target_cell = row.cells[0]
        source_cell = table.cell(0, 0)  # Use the first cell as a template

        # Copy cell properties from the template cell
        utils.copy_cell_properties(source_cell, target_cell)

        # Add paragraphs and runs to the new cell, copying formatting
        for p_i, paragraph in enumerate(paragraphs):
            if p_i == 0: 
                new_paragraph = target_cell.paragraphs[0]
            else: 
                new_paragraph = target_cell.add_paragraph()
            source_paragraph = source_cell.paragraphs[p_i]
            if p_i == 0:
                picture.add_float_picture(
                    new_paragraph,
                    "pictures/basic-cert-background.png",
                    width=CERT_WIDTH_INCHES,
                    height=CERT_HEIGHT_INCHES,
                )

            new_paragraph.alignment = source_paragraph.alignment
            new_paragraph.paragraph_format.left_indent = (
                source_paragraph.paragraph_format.left_indent
            )
            for target_run, source_run in zip(paragraph.runs, source_paragraph.runs):
                new_run = new_paragraph.add_run(target_run.text)
                utils.preserve_formatting(new_run, source_run)
    return final_doc


def add_student_content_to_merged_table(
    merged_table, tbl, student_index, curr_index, picture_path, picture_height=None, picture_width=None
):
    xml = etree.tostring(tbl._tbl, encoding='unicode', pretty_print=True)
    utils.dump(xml, 'og table')
    if student_index == 0:
        for element_name in ["w:tblGrid", "w:tblPr"]:
            utils.copy_table_element(tbl._tbl, merged_table._tbl, element_name)

    for row_index in range(
        len(tbl._tbl.findall("./w:tr", namespaces=tbl._tbl.nsmap))
    ):  # Iterate using XML
        target_row = merged_table.rows[curr_index]._element
        source_row_element = tbl.rows[row_index]._element

        utils.addTrPr(source_row_element, target_row)

        # --- Copy cells from source row to target row ---
        source_row_cells = source_row_element.findall(
            "./w:tc", namespaces=source_row_element.nsmap
        )
        for col_index, source_cell in enumerate(source_row_cells):
            target_cell = target_row[col_index]
            for child in source_cell:
                utils.update_nested_table_styles(source_cell, source_row_element)
                target_cell.append(copy.deepcopy(child))
        
        for cell in merged_table.rows[curr_index].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if "prof_educ_logo" in run.text:
                        run.text = run.text.replace("prof_educ_logo", "") 
                        run.add_picture('pictures/professional-education-logo.png')
                    if "bigger_educ_logo" in run.text:
                        run.text = run.text.replace("bigger_educ_logo", "") 
                        run.add_picture('pictures/professional-education-logo.png', width=Inches(1.53), height=Inches(1.09))
                    

        first_cell = merged_table.rows[curr_index].cells[0]
        curr_index += 1
        p = first_cell.add_paragraph()
        picture.add_float_picture(
            p,
            picture_path,
            height=picture_height,
            width=picture_width,
            pos_x=Pt(0),
            pos_y=Pt(0),
        )
    xml = etree.tostring(merged_table._tbl, encoding='unicode', pretty_print=True)
    utils.dump(xml, 'merged table')


def create_tractor_certificate(replacement_dict, students, picture_path):
    if not students:
        return Document()

    merged_doc = Document()
    merged_doc = utils.fit_more_rows(merged_doc)
    utils.set_default_font(merged_doc)

    merged_table = merged_doc.add_table(rows=len(students), cols=2)
    merged_tractor_table = merged_doc.add_table(rows=len(students), cols=2)

    curr_index = 0
    for student_index, student in enumerate(students):
        local_dict = replacement_dict.copy()
        local_dict[NAME_KEY] = student[NAME_KEY]
        local_dict[CERTIFICATE_KEY] = student[CERTIFICATE_KEY]
        local_dict[MACHINE_CATEGORY] = student[MACHINE_CATEGORY]

        doc = DocxTemplate("templates/certificate_tractor.docx")
        doc.render(local_dict)

        add_student_content_to_merged_table(
            merged_table, doc.tables[0], student_index, curr_index, picture_path, picture_height=TRACTOR_CERT_HEIGHT, picture_width=TRACTOR_CERT_WIDTH
        )
        add_student_content_to_merged_table(
            merged_tractor_table, doc.tables[1], student_index, curr_index, picture_path, picture_height=TRACTOR_CERT_HEIGHT, picture_width=TRACTOR_CERT_WIDTH
        )
        curr_index += 1

    return merged_doc


def create_tractor_certs(dict, students):
    blue = create_tractor_certificate(
        dict, students, "pictures/tractor-background-blue.png"
    )
    dict_with_profession_replaced = dict.copy()
    dict_with_profession_replaced["student_profession"] = TRACTOR_PROFESSION_WORDING
    green = create_tractor_certificate(
        dict_with_profession_replaced, students, "pictures/tractor-background-green.png"
    )
    return (blue, green)


def create_beginning_document(beginning_dict, students):
    """Creates a Word document with the provided information."""

    doc = DocxTemplate("templates/Приказ о начале.docx")
    doc.render(beginning_dict)
    utils.set_default_font(doc)
    table = doc.tables[0]

    for index, student in enumerate(students):
        # Create a new row
        new_row = table.add_row()

        # Populate cells in the new row
        new_row.cells[0].text = str(index + 1)
        new_row.cells[1].text = student[NAME_KEY]
        new_row.cells[2].text = beginning_dict["student_company"]

        # Add more cells for other data (profession, date, etc.) if needed
    return doc


def create_end_doc(replacement_dict, students):
    """Creates a Word document with the provided information."""

    doc = DocxTemplate("templates/Приказ о выпуске.docx")
    doc.render(replacement_dict)
    utils.set_default_font(doc)
    table = doc.tables[0]

    for index, student in enumerate(students):
        # Create a new row
        new_row = table.add_row()

        # Populate cells in the new row
        new_row.cells[0].text = str(index + 1)
        new_row.cells[1].text = student[NAME_KEY]
        new_row.cells[2].text = replacement_dict["student_company"]
        new_row.cells[3].text = str(student[CERTIFICATE_KEY])

    return doc


def create_protocol_doc(replacement_dict, students):
    """Creates a Word document with the provided information."""

    doc = DocxTemplate("templates/Протокол.docx")
    doc.render(replacement_dict)
    utils.set_default_font(doc)
    table = doc.tables[0]

    for index, student in enumerate(students):
        # Create a new row
        new_row = table.add_row()

        # Populate cells in the new row
        new_row.cells[0].text = str(index + 1)
        new_row.cells[1].text = student[NAME_KEY]
        new_row.cells[2].text = replacement_dict["student_company"]
        new_row.cells[3].text = str(student[CERTIFICATE_KEY])

    return doc


st.title("Профессиональное обучение")

# Input 1: Text Input
available_professions = utils.load_from_pickle("data/professions.pickle")
student_profession = utils.choose_profession(available_professions)

today = datetime.date.today()
beginning_date = st.date_input("дата начала", value=today)
end_date = st.date_input("дата окончания", value=today)

beginning_number = st.number_input(
    "номер приказа о начале", step=1, value=1, placeholder=808
)
end_number = st.number_input(
    "номер приказа об окончании", step=1, value=1, placeholder=808
)

# this should be replaced by a scroll through
teacher_name = utils.choose_teacher(utils.load_from_pickle("data/teachers.pickle"))

company = st.text_input(
    "Предприятие", "заявление", placeholder="Наименование предприятия или 'заявление'"
)
student_names = st.text_area("Введите имена студентов, по одному на строку").split("\n")
student_data = []
for line in student_names:
    if line:
        items = [item for item in line.split("\t") if item]
        certificate_number, _, _, name, *category = (
            items  # Split each line by tab
        )
        student_data.append(
            {
                NAME_KEY: name,
                CERTIFICATE_KEY: int(float(certificate_number)),
                MACHINE_CATEGORY: category[0] if category else "",
            }
        )
num_students = len(student_data)

formatted_beginning_date = utils.format_date(beginning_date)
formatted_end_date = utils.format_date(end_date)
replacement_dict = {
    "beginning_date": formatted_beginning_date,
    "beginning_number": beginning_number,
    "end_date": formatted_end_date,
    "end_number": end_number,
    "student_company": company,
    "teacher_name": teacher_name,
    "num_students": num_students,
    "class": "4",
    "year": beginning_date.year
}
if student_profession: 
    if student_profession.hours_str:
        replacement_dict['hours'] = student_profession.hours_str
    if student_profession.formatted_profession: 
        replacement_dict['student_profession'] = student_profession.formatted_profession

beginning_doc = create_beginning_document(replacement_dict, student_data)
end_doc = create_end_doc(replacement_dict, student_data)
protocol = create_protocol_doc(replacement_dict, student_data)
certificate_docs = create_certificate(replacement_dict, student_data)
(blue_tractor_cert, green_tractor_cert) = create_tractor_certs(
    replacement_dict, student_data
)
milana_conf_page = create_confirmation_page(replacement_dict, student_data, 'pictures/tractor-background-green.png')
milana_cert = create_certificate_for_labour_protection(replacement_dict, student_data)

show_documents = st.button("Сгенерировать документы")

if show_documents:
    if not student_profession:
        st.warning("Укажите профессию")
    if not teacher_name:
        st.warning("Укажите преподователя")
    if not beginning_date:
        st.warning("Укажите дату начала")
    if not end_date:
        st.warning("Укажите дату окончания")
    if not beginning_number:
        st.warning("Укажите номер приказа о начале")
    if not end_number:
        st.warning("Укажите номер приказа о выпуске")
    if num_students == 0:
        st.warning("Укажите обучающихся")

    if (
        student_profession
        and teacher_name
        and beginning_date
        and end_date
        and beginning_number
        and end_number
    ):
        document_tabs = st.tabs(
            [
                "Приказ о начале",
                "Приказ об окончании",
                "Протокол",
                "Свидетельство",
                "Свидетельство тракторов синее",
                "Свидетельство тракторов зеленое",
                "Милана удостоверение", 
                "Милана св-во охрана труда",
            ]
        )
        with document_tabs[0]:  # Приказ о начале
            utils.display_docx_content(beginning_doc)

        with document_tabs[1]:  # Приказ об окончании
            utils.display_docx_content(end_doc)

        with document_tabs[2]:  # Протокол
            utils.display_docx_content(protocol)

        with document_tabs[3]:  # Свидетельство
            utils.display_docx_content(certificate_docs)

        with document_tabs[4]:  # Свидетельство тракторов
            utils.display_docx_content(blue_tractor_cert)

        with document_tabs[5]:
            utils.display_docx_content(green_tractor_cert)
        with document_tabs[6]: 
            utils.display_docx_content(milana_conf_page)
        with document_tabs[7]: 
            utils.display_docx_content(milana_cert)

# --- Create a ZIP archive in memory ---
zip_buffer = BytesIO()
with zipfile.ZipFile(zip_buffer, "w") as zipf:
    # Add the beginning document
    with zipf.open("Приказ о начале.docx", "w") as f:
        beginning_doc.save(f)

    # Add the end document
    with zipf.open("Приказ о выпуске.docx", "w") as f:
        end_doc.save(f)

    with zipf.open("Протокол.docx", "w") as f:
        protocol.save(f)

    with zipf.open("Свидетельство.docx", "w") as f:
        certificate_docs.save(f)

    with zipf.open("Свидетельство синее трактор.docx", "w") as f:
        blue_tractor_cert.save(f)

    with zipf.open("Свидетельство зеленое трактор.docx", "w") as f:
        green_tractor_cert.save(f)
    with zipf.open("Удостоверение Милана.docx", "w") as f:
        milana_conf_page.save(f)
    with zipf.open("Свидетельство Милана.docx", "w") as f:
        milana_cert.save(f)

zip_buffer.seek(0)

formatted_end_date = end_date.strftime("%d.%m.%Y")
# --- Download the ZIP archive ---
st.download_button(
    label="Скачать документы (ZIP)",
    data=zip_buffer,
    file_name=f"{formatted_end_date}.zip",
    mime="application/zip",
)
