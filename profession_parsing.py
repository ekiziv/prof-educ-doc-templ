from docx import Document
import streamlit as st
import pickle
import utils 


def format_profession_string(profession_name, codes):
    if len(codes) > 0:
        code_str = ", ".join(
            str(code) for code in codes
        )
        return f"{code_str} «{profession_name}»"
    return f"«{profession_name}»"

# return a list of ints representing the applicable codes 
def add_code(code_str): 
    code_list = []
    if code_str == None: 
        return code_list
    
    try:
        code_list.append(int(code_str))
    except ValueError:
        if code_str == "-":
            return code_list
        else:
            # Handle cases with multiple codes (assuming comma-separated)
            codes = [int(c.strip()) for c in code_str.split(",") if c.strip()]
            for code in codes: 
                code_list.append(code)
    return code_list

def add_hours(hours_str): 
    if hours_str == None: 
        return None
    
    try:
        hours_int = int(hours_str)
        return utils.format_hours_string(hours_int)
    except ValueError:
        print('Failed to parse hours')
        return None

def professions_docx_table_to_df(docx_path='/Users/ekiziv/Desktop/mama/work/data/all_professions_cleaned.docx'):
    """Loads a table from a .docx file into a dictionary and saves to pickle. 
    Handles multiple codes per profession.

    Args:
        docx_path (str, optional): Path to the .docx file. 
    Returns:
        dict: Dictionary containing the table data (profession: [list of codes]).
    """

    doc = Document(docx_path)
    table = doc.tables[0]
    professions = {}

    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        profession = row_data[0]
        code = row_data[1]

        if not profession:  # Skip empty rows
            continue
        
        code_list = add_code(code)
        
        # No hours specified for professions from this list.
        new_profession = utils.Profession(name=profession, code=code_list, hours_str=None, formatted_profession=format_profession_string(profession, code_list), role_required=False)
        professions[profession] = new_profession

    return professions

def professions_labour_protection(docx_path='/Users/ekiziv/Desktop/mama/work/data/milana_professions.docx'):
    doc = Document(docx_path)
    table = doc.tables[0]
    professions = {}

    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        profession = row_data[0]
        hours = row_data[1]

        if not profession or not hours:  # Skip rows with either no profession or hours
            continue

        hrs = add_hours(hours)
        
        code_list = []
        new_profession = utils.Profession(name=profession, code=code_list, hours_str=hrs, formatted_profession=format_profession_string(profession, code_list), role_required=True)
        professions[profession] = new_profession

    return professions

def teachers_to_pickle():
    teachers = ["А.И. Мамонтов", 
                "А.В. Перекрестов", 
                "Н.В. Клюшина",
                "Л.А. Лапчук"]
    with open("data/teachers.pickle", "wb") as f:
        pickle.dump(teachers, f)

if __name__ == "__main__":
    professions = professions_docx_table_to_df()
    hourly_professions = professions_labour_protection()
    professions.update(hourly_professions)
    professions = dict(sorted(professions.items()))
    st.write(professions)

    with open("data/professions.pickle", "wb") as f:
        pickle.dump(professions, f)

    teachers_to_pickle()